from docx import Document
from docx.shared import Inches
import pandas as pd
from docx2pdf import convert
import os


def convert_docx_to_pdf(docx_path, pdf_path):
    convert(docx_path, pdf_path)


def replace_placeholder_in_docx(doc_path, placeholder, replacement_text, output_path,img_path):
    # Load the DOCX document
    doc = Document(doc_path)
    
    # Überprüfen und Ersetzen von Platzhaltern in Tabellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    #for i in range(len(placeholder)-1):
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            if placeholder in run.text: 
                                run.text = run.text.replace(placeholder, replacement_text)
                                if img_path != '':
                                    run.add_picture(img_path, width=Inches(3.75))
    
    # Save the modified document
    doc.save(output_path)
    print(f"Document saved as: {output_path}")

#create the weekly report
def ToWord(zahl, h_id, doc_path, week):
    id=h_id
    replacing = pd.read_excel('../'+str(id)+'/Week_'+str(week)+'/placeholder_week_'+str(week)+'_id_'+str(id)+'.xlsx')

    output_path = doc_path
    
    placeholder = replacing['Placeholder'][zahl]

    if 'Image' in placeholder:
        img_path = '../'+str(id)+'/Week_'+str(week)+'/'+replacing['Content'][zahl]
        if os.path.exists(img_path):
            replacement_text = ''
        else:
            img_path =''
            replacement_text = 'Es gibt dieses mal keine Daten für diese Grafik.'
    else:
        if isinstance(replacing['Content'][zahl], float):
            replacement_text = ''
            print('Detected')
        else:
            replacement_text = replacing['Content'][zahl]
        img_path =''

    replace_placeholder_in_docx(doc_path, placeholder, replacement_text, output_path,img_path)
    
# create the monthly report
def ToWordMonth(zahl, h_id, doc_path, month):
    id=h_id
    replacing = pd.read_excel('../'+str(id)+'/Month_'+str(month)+'/placeholder_month_'+str(month)+'_id_'+str(id)+'.xlsx')

    output_path = doc_path
    
    placeholder = replacing['Placeholder'][zahl]

    if 'Image' in placeholder:
        img_path = '../'+str(id)+'/Month_'+str(month)+'/'+replacing['Content'][zahl]
        if os.path.exists(img_path):
            replacement_text = ''
        else:
            img_path =''
            replacement_text = ''
    else: 
        if isinstance(replacing['Content'][zahl], float):
            replacement_text = ''
            print('Detected')
        else:
            replacement_text = replacing['Content'][zahl]
        img_path =''

    replace_placeholder_in_docx(doc_path, placeholder, replacement_text, output_path,img_path)

#################################
#################################
#################################
# Example Data
id = 1758394 # this is hotel_id
week = 46
month = 11

#for weekly report
#load the data for generating the reports from the respective excel files
df = pd.read_excel('../'+str(id)+'/Week_'+str(week)+'/placeholder_week_'+str(week)+'_id_'+str(id)+'.xlsx')
doc_path = '../'+str(id)+'/Week_'+str(week)+'/report_week_'+str(week)+'_id_'+str(id)+'.docx'
output_path = '../'+str(id)+'/Week_'+str(week)+'/report_week_'+str(week)+'_id_'+str(id)+'.pdf'

#execute the generation of the report
for i in range(len(df['Placeholder'])):
    ToWord(i,id, doc_path, week)

convert_docx_to_pdf(doc_path, output_path)

##########

# For the monthly report
#load the data for generating the reports from the respective excel files
df = pd.read_excel('../'+str(id)+'/Month_'+str(month)+'/placeholder_month_'+str(month)+'_id_'+str(id)+'.xlsx')
doc_path = '../'+str(id)+'/Month_'+str(month)+'/report_month_'+str(month)+'_id_'+str(id)+'.docx'
output_path = '../'+str(id)+'/Month_'+str(month)+'/report_month_'+str(month)+'_id_'+str(id)+'.pdf'

#execute the generation of the report
for i in range(len(df['Placeholder'])):
    ToWordMonth(i,id, doc_path, month)

convert_docx_to_pdf(doc_path, output_path)
